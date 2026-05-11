"""Génère le guide de déploiement Word pour le projet ML Housing sur Render.com."""

import datetime
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

doc = Document()

# ── Styles ────────────────────────────────────────────────────────────────────


def set_font(run, bold=False, italic=False, size=None, color=None, mono=False):
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if mono:
        run.font.name = "Courier New"


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.runs[0].font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    return h


def add_info_box(doc, text, color=(0xEF, 0xF6, 0xFF), border_color=(0x3B, 0x82, 0xF6)):
    """Encadré bleu informatif."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "{:02X}{:02X}{:02X}".format(*color))
    shading.set(qn("w:val"), "clear")
    p._p.get_or_add_pPr().append(shading)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    return p


def add_warning_box(doc, text):
    """Encadré orange avertissement."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "FFF7ED")
    shading.set(qn("w:val"), "clear")
    p._p.get_or_add_pPr().append(shading)
    run = p.add_run("⚠  " + text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xC2, 0x41, 0x0C)
    return p


def add_success_box(doc, text):
    """Encadré vert succès."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F0FDF4")
    shading.set(qn("w:val"), "clear")
    p._p.get_or_add_pPr().append(shading)
    run = p.add_run("✔  " + text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x16, 0x6D, 0x3B)
    return p


def add_code(doc, code, title=None):
    """Bloc de code monospace avec fond gris."""
    if title:
        t = doc.add_paragraph()
        tr = t.add_run(title)
        tr.bold = True
        tr.font.size = Pt(9)
        tr.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F3F4F6")
    shading.set(qn("w:val"), "clear")
    p._p.get_or_add_pPr().append(shading)
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    return p


def add_step(doc, number, title, description=None):
    """Étape numérotée mise en valeur."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r1 = p.add_run(f"  Étape {number} — ")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    r2 = p.add_run(title)
    r2.bold = True
    r2.font.size = Pt(11)
    if description:
        doc.add_paragraph(description).runs[0].font.size = (
            Pt(10) if doc.paragraphs[-1].runs else None
        )
    return p


def add_table_kv(doc, rows, col1="Variable", col2="Valeur / Exemple"):
    """Tableau à deux colonnes clé/valeur."""
    table = doc.add_table(rows=len(rows) + 1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = col1
    hdr[1].text = col2
    for cell in hdr:
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "1A56DB")
        shading.set(qn("w:val"), "clear")
        cell._tc.get_or_add_tcPr().append(shading)
    for i, (k, v) in enumerate(rows, 1):
        cells = table.rows[i].cells
        cells[0].text = k
        cells[1].text = v
        if i % 2 == 0:
            for cell in cells:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "F9FAFB")
                shading.set(qn("w:val"), "clear")
                cell._tc.get_or_add_tcPr().append(shading)
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    return table


# ── Page de titre ─────────────────────────────────────────────────────────────

section = doc.sections[0]
section.page_width = Inches(8.27)
section.page_height = Inches(11.69)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1.18)
section.right_margin = Inches(1.18)

title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_para.paragraph_format.space_before = Pt(60)
r = title_para.add_run("Guide de Déploiement")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

sub_para = doc.add_paragraph()
sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub_para.add_run("ML Housing Prediction — Mise en production sur Render.com")
r2.font.size = Pt(15)
r2.font.color.rgb = RGBColor(0x37, 0x4B, 0x73)

doc.add_paragraph()
date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = date_para.add_run(f"Version 1.0  •  {datetime.date.today().strftime('%d %B %Y')}")
r3.font.size = Pt(10)
r3.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
r3.italic = True

doc.add_page_break()

# ── Sommaire ──────────────────────────────────────────────────────────────────

add_heading(doc, "Sommaire", 1)
toc_items = [
    ("1", "Prérequis et architecture", 3),
    ("2", "Partie 1 — Stockage des modèles avec Cloudflare R2", 3),
    ("3", "Partie 2 — Configuration du compte Render", 3),
    ("4", "Partie 3 — Déploiement du service Backend", 3),
    ("5", "Partie 4 — Déploiement du service Frontend", 3),
    ("6", "Partie 5 — Configuration GitHub pour le CI/CD", 3),
    ("7", "Partie 6 — Premier déploiement via le workflow", 3),
    ("8", "Partie 7 — Vérification et tests", 3),
    ("9", "Partie 8 — Dépannage", 3),
    ("10", "Annexes — Récapitulatif complet des variables", 3),
]
for num, title, _ in toc_items:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(f"{num}.  {title}")
    r.font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — PRÉREQUIS
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "1. Prérequis et architecture", 1)

add_heading(doc, "1.1  Ce dont vous avez besoin avant de commencer", 2)

doc.add_paragraph(
    "Avant de suivre ce guide, assurez-vous de disposer des éléments suivants :"
)
prereqs = [
    "Un compte GitHub avec accès au dépôt rahiming/ml-housting-project",
    "Le modèle ML entraîné localement : artifacts/models/model_latest.joblib\n"
    "   → Si absent : exécuter le pipeline d'entraînement (python -m src.training.train)",
    "Un navigateur web moderne (Chrome, Firefox, Edge)",
    "Accès à un terminal avec Python 3.10+ et pip installés",
    "Une adresse e-mail valide (pour créer les comptes Cloudflare et Render)",
]
for item in prereqs:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(item).font.size = Pt(10)

add_heading(doc, "1.2  Architecture de l'application déployée", 2)

doc.add_paragraph(
    "L'application se compose de trois éléments qui communiquent entre eux :"
)
add_code(
    doc,
    """\
┌─────────────────────────────────────────────────────────────────────┐
│                        INTERNET (HTTPS)                             │
│                                                                     │
│   Utilisateur                                                       │
│       │                                                             │
│       ▼                                                             │
│  ┌────────────────────┐         ┌────────────────────────────┐      │
│  │  Frontend Streamlit │ ──────► │   Backend FastAPI           │      │
│  │  (Render – port 8501)│  HTTP  │   (Render – port 8000)     │      │
│  └────────────────────┘         └──────────────┬─────────────┘      │
│                                                │                    │
│                                                ▼                    │
│                                  ┌─────────────────────────┐        │
│                                  │  Cloudflare R2 (S3)     │        │
│                                  │  Stockage des modèles   │        │
│                                  │  model_latest.joblib    │        │
│                                  │  model_v1.joblib        │        │
│                                  │  model_v2.joblib        │        │
│                                  └─────────────────────────┘        │
└─────────────────────────────────────────────────────────────────────┘\
""",
)

add_info_box(
    doc,
    "Render est une plateforme cloud qui héberge des conteneurs Docker directement "
    "depuis un registre d'images (ici GHCR — GitHub Container Registry). "
    "Cloudflare R2 remplace MinIO comme stockage objet S3-compatible, "
    "avec un tier gratuit généreux (10 Go / mois, 0 frais de sortie).",
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PARTIE 1 — CLOUDFLARE R2
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "2. Partie 1 — Stockage des modèles avec Cloudflare R2", 1)
doc.add_paragraph(
    "Le backend a besoin de télécharger les fichiers de modèles (.joblib) au démarrage. "
    "Sur Render il n'y a pas de système de fichiers persistant entre les démarrages, "
    "donc les modèles doivent être stockés dans un service S3 externe. "
    "Cloudflare R2 est l'option la plus simple : gratuit jusqu'à 10 Go, "
    "API 100 % compatible AWS S3, aucun frais de bande passante sortante."
)

add_heading(doc, "2.1  Création du compte Cloudflare", 2)

add_step(doc, 1, "Ouvrir le site de Cloudflare")
doc.add_paragraph("Naviguer vers : https://dash.cloudflare.com/sign-up")
p = doc.add_paragraph(style="List Bullet")
p.add_run("Renseigner l'adresse e-mail et un mot de passe fort.").font.size = Pt(10)
p = doc.add_paragraph(style="List Bullet")
p.add_run("Cliquer sur « Create Account ».").font.size = Pt(10)
p = doc.add_paragraph(style="List Bullet")
p.add_run("Confirmer l'e-mail via le lien reçu.").font.size = Pt(10)

add_warning_box(
    doc,
    "Si vous avez déjà un compte Cloudflare (même personnel), vous pouvez l'utiliser. "
    "R2 est disponible sur tous les comptes y compris le plan Free.",
)

add_heading(doc, "2.2  Activation de Cloudflare R2", 2)

add_step(doc, 2, "Accéder à R2 depuis le dashboard Cloudflare")
doc.add_paragraph(
    "Une fois connecté au dashboard Cloudflare (https://dash.cloudflare.com) :"
)
p = doc.add_paragraph(style="List Number")
p.add_run("Dans le menu de gauche, cliquer sur « R2 Object Storage ».").font.size = Pt(
    10
)
p = doc.add_paragraph(style="List Number")
p.add_run(
    "Si c'est la première utilisation, cliquer sur « Purchase R2 Plan » "
    "→ choisir le plan Free (pas de carte requise pour le tier gratuit)."
).font.size = Pt(10)
p = doc.add_paragraph(style="List Number")
p.add_run("Attendre que le service soit activé (quelques secondes).").font.size = Pt(10)

add_heading(doc, "2.3  Création du bucket", 2)

add_step(doc, 3, "Créer le bucket qui contiendra les modèles ML")
p = doc.add_paragraph(style="List Number")
p.add_run("Cliquer sur « Create bucket ».").font.size = Pt(10)
p = doc.add_paragraph(style="List Number")
r = p.add_run("Bucket name : ")
r.font.size = Pt(10)
r.bold = True
r2 = p.add_run("ml-models")
r2.font.name = "Courier New"
r2.font.size = Pt(10)
p = doc.add_paragraph(style="List Number")
p.add_run(
    "Location : choisir « Automatic » (Cloudflare sélectionne la région optimale)."
).font.size = Pt(10)
p = doc.add_paragraph(style="List Number")
p.add_run("Cliquer sur « Create bucket ».").font.size = Pt(10)

add_success_box(
    doc, "Le bucket ml-models est créé. Vous êtes redirigé vers sa page de gestion."
)

add_heading(doc, "2.4  Récupération de l'Account ID Cloudflare", 2)

add_step(doc, 4, "Trouver l'Account ID (nécessaire pour construire l'endpoint S3)")
doc.add_paragraph(
    "L'endpoint S3 de Cloudflare R2 suit le format :\n"
    "https://<ACCOUNT_ID>.r2.cloudflarestorage.com"
)
p = doc.add_paragraph(style="List Number")
p.add_run(
    "Dans le menu de gauche, cliquer sur le nom de votre compte "
    "(en haut à gauche) ou aller dans « Overview »."
).font.size = Pt(10)
p = doc.add_paragraph(style="List Number")
p.add_run(
    "L'Account ID est affiché dans le panneau de droite (section « API »). "
    "Il ressemble à : a1b2c3d4e5f6789012345678901234ab"
).font.size = Pt(10)
p = doc.add_paragraph(style="List Number")
p.add_run("Copier et conserver cet ID.").font.size = Pt(10)

add_code(
    doc,
    "MINIO_ENDPOINT = https://<ACCOUNT_ID>.r2.cloudflarestorage.com\n"
    "Exemple réel  : https://a1b2c3d4e5f6789012345678901234ab.r2.cloudflarestorage.com",
    title="Construction de la variable MINIO_ENDPOINT",
)

add_heading(doc, "2.5  Création des tokens d'accès R2", 2)

add_step(doc, 5, "Générer les clés d'accès API pour le bucket")
p = doc.add_paragraph(style="List Number")
p.add_run(
    "Dans R2, cliquer sur « Manage R2 API Tokens » (lien en haut à droite de la page R2)."
).font.size = Pt(10)
p = doc.add_paragraph(style="List Number")
p.add_run("Cliquer sur « Create API token ».").font.size = Pt(10)
p = doc.add_paragraph(style="List Number")
r = p.add_run("Token name : ")
r.font.size = Pt(10)
r.bold = True
p.add_run("ml-housing-backend-rw").font.size = Pt(10)
p = doc.add_paragraph(style="List Number")
p.add_run("Permissions : sélectionner « Object Read & Write ».").font.size = Pt(10)
p = doc.add_paragraph(style="List Number")
p.add_run(
    "Specify bucket (optionnel mais recommandé) : sélectionner « ml-models »."
).font.size = Pt(10)
p = doc.add_paragraph(style="List Number")
p.add_run("Cliquer sur « Create API Token ».").font.size = Pt(10)

add_warning_box(
    doc,
    "Cloudflare n'affiche la Secret Access Key qu'UNE SEULE FOIS. "
    "Copier immédiatement les deux valeurs suivantes et les conserver en sécurité :",
)

add_table_kv(
    doc,
    [
        ("Access Key ID", "Ressemble à : a1b2c3d4e5f6g7h8i9j0k1l2m3n4o5p6"),
        ("Secret Access Key", "Ressemble à : Aa1Bb2Cc3Dd4Ee5Ff6Gg7Hh8Ii9Jj0Kk1Ll2Mm3"),
    ],
    col1="Clé",
    col2="Format",
)

add_heading(doc, "2.6  Upload des fichiers de modèles", 2)

add_step(doc, 6, "Uploader les modèles ML dans le bucket R2")
doc.add_paragraph(
    "Trois fichiers de modèles doivent être uploadés dans le bucket ml-models. "
    "Ils se trouvent localement dans le dossier artifacts/models/ de votre projet."
)
add_warning_box(
    doc,
    "Si les fichiers model_v1.joblib et model_v2.joblib n'existent pas encore, "
    "exécuter d'abord le script scripts/upload_model_to_minio.py en local "
    "avec un MinIO local, ou entraîner et exporter les modèles A/B manuellement.",
)

doc.add_paragraph("Option A — Via l'interface web Cloudflare R2 :")
p = doc.add_paragraph(style="List Number")
p.add_run("Dans le dashboard R2, ouvrir le bucket ml-models.").font.size = Pt(10)
p = doc.add_paragraph(style="List Number")
p.add_run("Cliquer sur « Upload » (bouton en haut à droite).").font.size = Pt(10)
p = doc.add_paragraph(style="List Number")
p.add_run("Glisser-déposer ou sélectionner les fichiers :").font.size = Pt(10)
for fname in ["model_latest.joblib", "model_v1.joblib", "model_v2.joblib"]:
    p2 = doc.add_paragraph(style="List Bullet 2")
    r = p2.add_run(fname)
    r.font.name = "Courier New"
    r.font.size = Pt(10)
p = doc.add_paragraph(style="List Number")
p.add_run("Cliquer sur « Upload » pour confirmer.").font.size = Pt(10)

doc.add_paragraph("Option B — Via rclone (ligne de commande) :")
add_code(
    doc,
    "# Installation de rclone\n"
    "pip install rclone  # ou télécharger depuis rclone.org\n\n"
    "# Configuration rclone pour Cloudflare R2\n"
    "rclone config create r2 s3 \\\n"
    "  provider Cloudflare \\\n"
    "  access_key_id <VOTRE_ACCESS_KEY_ID> \\\n"
    "  secret_access_key <VOTRE_SECRET_ACCESS_KEY> \\\n"
    "  endpoint https://<ACCOUNT_ID>.r2.cloudflarestorage.com\n\n"
    "# Upload des modèles\n"
    "rclone copy artifacts/models/ r2:ml-models/",
    title="Upload via rclone",
)

add_success_box(
    doc,
    "Vérification : dans le dashboard R2, le bucket ml-models doit contenir "
    "model_latest.joblib, model_v1.joblib et model_v2.joblib.",
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PARTIE 2 — COMPTE RENDER
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "3. Partie 2 — Configuration du compte Render", 1)

add_heading(doc, "3.1  Création du compte Render", 2)

add_step(doc, 7, "Créer un compte sur render.com")
p = doc.add_paragraph(style="List Number")
p.add_run(
    "Naviguer vers https://render.com et cliquer sur « Get Started for Free »."
).font.size = Pt(10)
p = doc.add_paragraph(style="List Number")
p.add_run(
    "Choisir « Sign up with GitHub » pour lier directement le compte GitHub "
    "(recommandé : simplifie l'authentification)."
).font.size = Pt(10)
p = doc.add_paragraph(style="List Number")
p.add_run("Autoriser Render à accéder à votre compte GitHub.").font.size = Pt(10)
p = doc.add_paragraph(style="List Number")
p.add_run("Confirmer l'e-mail si demandé.").font.size = Pt(10)

add_heading(doc, "3.2  Configuration des Registry Credentials (GHCR)", 2)

doc.add_paragraph(
    "Les images Docker du projet sont stockées dans GHCR (GitHub Container Registry) "
    "sous l'adresse ghcr.io/rahiming/ml-housting-project/. "
    "Render a besoin d'un credential pour pouvoir puller ces images privées."
)

add_step(doc, 8, "Créer un Personal Access Token GitHub pour Render")
doc.add_paragraph(
    "Ce token permettra à Render de télécharger vos images Docker depuis GHCR."
)
p = doc.add_paragraph(style="List Number")
p.add_run("Aller sur https://github.com/settings/tokens/new").font.size = Pt(10)
p = doc.add_paragraph(style="List Number")
r = p.add_run("Note : ")
r.font.size = Pt(10)
r.bold = True
p.add_run("render-ghcr-pull").font.size = Pt(10)
p = doc.add_paragraph(style="List Number")
p.add_run(
    "Expiration : 90 days (ou No expiration pour un projet de cours)."
).font.size = Pt(10)
p = doc.add_paragraph(style="List Number")
p.add_run("Scopes : cocher uniquement ").font.size = Pt(10)
p.add_run("read:packages").font.name = "Courier New"
p.runs[-1].font.size = Pt(10)
p = doc.add_paragraph(style="List Number")
p.add_run(
    "Cliquer sur « Generate token » et copier le token (ghp_xxxx…)."
).font.size = Pt(10)

add_warning_box(doc, "Ce token n'est visible qu'une seule fois. Le copier maintenant.")

add_step(doc, 9, "Enregistrer le credential dans Render")
p = doc.add_paragraph(style="List Number")
p.add_run(
    "Dans Render, cliquer sur l'icône de profil en haut à droite "
    "→ « Account Settings »."
).font.size = Pt(10)
p = doc.add_paragraph(style="List Number")
p.add_run("Dans le menu de gauche, cliquer sur « Registry Credentials ».").font.size = (
    Pt(10)
)
p = doc.add_paragraph(style="List Number")
p.add_run("Cliquer sur « New Credential ».").font.size = Pt(10)

add_table_kv(
    doc,
    [
        ("Name", "ghcr-ml-housing"),
        ("Registry", "GitHub Container Registry"),
        ("Username", "rahiming  (votre GitHub username)"),
        ("Password", "ghp_xxxx… (le token créé à l'étape 8)"),
    ],
    col1="Champ",
    col2="Valeur",
)

p = doc.add_paragraph(style="List Number")
p.add_run("Cliquer sur « Save Credential ».").font.size = Pt(10)

add_heading(doc, "3.3  Génération de la clé API Render", 2)

add_step(doc, 10, "Créer une clé API Render pour le pipeline CI/CD")
doc.add_paragraph(
    "Le workflow GitHub Actions a besoin d'une clé API Render pour déclencher "
    "les déploiements et vérifier leur statut."
)
p = doc.add_paragraph(style="List Number")
p.add_run("Dans Render, aller dans « Account Settings » → « API Keys ».").font.size = (
    Pt(10)
)
p = doc.add_paragraph(style="List Number")
p.add_run("Cliquer sur « Create API Key ».").font.size = Pt(10)
r = p.add_run("")
p = doc.add_paragraph(style="List Number")
r = p.add_run("Name : ")
r.font.size = Pt(10)
r.bold = True
p.add_run("github-actions-deploy").font.size = Pt(10)
p = doc.add_paragraph(style="List Number")
p.add_run("Cliquer sur « Create API Key » et copier la clé (rnd_xxxx…).").font.size = (
    Pt(10)
)

add_warning_box(doc, "Cette clé n'est visible qu'une seule fois. La copier maintenant.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PARTIE 3 — SERVICE BACKEND
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "4. Partie 3 — Déploiement du service Backend", 1)

add_heading(doc, "4.1  Création du service FastAPI Backend", 2)

add_step(doc, 11, "Créer un nouveau Web Service dans Render")
p = doc.add_paragraph(style="List Number")
p.add_run(
    "Dans le dashboard Render, cliquer sur le bouton « + New » (en haut à droite)."
).font.size = Pt(10)
p = doc.add_paragraph(style="List Number")
p.add_run("Sélectionner « Web Service ».").font.size = Pt(10)
p = doc.add_paragraph(style="List Number")
p.add_run("Sur la page suivante, choisir « Deploy an existing image ».").font.size = Pt(
    10
)

add_step(doc, 12, "Configurer l'image Docker du backend")

add_table_kv(
    doc,
    [
        ("Image URL", "ghcr.io/rahiming/ml-housting-project/backend:latest"),
        ("Registry Credential", "ghcr-ml-housing  (créé à l'étape 9)"),
    ],
    col1="Champ",
    col2="Valeur",
)

p = doc.add_paragraph(style="List Number")
p.add_run("Cliquer sur « Next ».").font.size = Pt(10)

add_step(doc, 13, "Renseigner les paramètres du service backend")

add_table_kv(
    doc,
    [
        ("Name", "ml-housing-backend"),
        ("Region", "Frankfurt (EU Central) — ou la région la plus proche"),
        (
            "Instance Type",
            "Free (512 MB RAM, 0.1 CPU) pour un projet de cours\n"
            "Starter (512 MB RAM, 0.5 CPU) pour de meilleures performances",
        ),
        ("Port", "8000"),
    ],
    col1="Champ",
    col2="Valeur",
)

add_warning_box(
    doc,
    "Le plan Free de Render met les services en veille après 15 minutes d'inactivité. "
    "Le premier appel après une période d'inactivité peut prendre 30 à 60 secondes "
    "(cold start). Pour un projet de démonstration c'est acceptable.",
)

add_step(doc, 14, "Ajouter les variables d'environnement du backend")
doc.add_paragraph(
    "Dans la section « Environment Variables » de la page de création du service, "
    "ajouter les variables suivantes. "
    "Les variables marquées « Secret » seront masquées dans les logs Render."
)

add_table_kv(
    doc,
    [
        (
            "MINIO_ENDPOINT",
            "https://<ACCOUNT_ID>.r2.cloudflarestorage.com\n"
            "Remplacer <ACCOUNT_ID> par votre Account ID Cloudflare",
        ),
        ("MINIO_BUCKET_MODELS", "ml-models"),
        ("MODEL_OBJECT_NAME", "model_latest.joblib"),
        ("MODEL_A_OBJECT_NAME", "model_v1.joblib"),
        ("MODEL_B_OBJECT_NAME", "model_v2.joblib"),
        ("AB_TRAFFIC_B_PERCENT", "50"),
        ("MINIO_ACCESS_KEY", "← coller l'Access Key ID R2  (marquer comme Secret)"),
        (
            "MINIO_SECRET_KEY",
            "← coller la Secret Access Key R2  (marquer comme Secret)",
        ),
    ],
    col1="Variable",
    col2="Valeur",
)

add_info_box(
    doc,
    "Pour marquer une variable comme Secret dans Render : cliquer sur l'icône "
    "en forme d'œil à côté du champ de valeur. La valeur sera chiffrée et "
    "masquée dans les logs et l'interface.",
)

add_step(doc, 15, "Déployer le service backend")
p = doc.add_paragraph(style="List Number")
p.add_run("Cliquer sur « Deploy Web Service » (bouton en bas de page).").font.size = Pt(
    10
)
p = doc.add_paragraph(style="List Number")
p.add_run(
    "Render va télécharger l'image Docker depuis GHCR et démarrer le conteneur. "
    "Cette opération prend généralement 2 à 5 minutes."
).font.size = Pt(10)
p = doc.add_paragraph(style="List Number")
p.add_run("Surveiller les logs dans l'onglet « Logs » du service.").font.size = Pt(10)
p = doc.add_paragraph(style="List Number")
p.add_run(
    "Attendre le message : Application prete a accepter des requetes."
).font.size = Pt(10)

add_heading(doc, "4.2  Récupération du Service ID et de l'URL Backend", 2)

add_step(doc, 16, "Noter le Service ID et l'URL du backend")
p = doc.add_paragraph(style="List Number")
p.add_run("Dans le service backend, aller dans l'onglet « Settings ».").font.size = Pt(
    10
)
p = doc.add_paragraph(style="List Number")
p.add_run(
    "Copier la valeur du champ « Service ID » (format : srv-xxxxxxxxxxxxxx). "
    "Cette valeur sera utilisée plus tard dans GitHub."
).font.size = Pt(10)
p = doc.add_paragraph(style="List Number")
p.add_run(
    "En haut de la page du service, noter l'URL publique du backend "
    "(format : https://ml-housing-backend.onrender.com)."
).font.size = Pt(10)

add_code(
    doc,
    "Service ID  : srv-xxxxxxxxxxxxxx   ← à copier dans GitHub (variable)\n"
    "URL backend : https://ml-housing-backend.onrender.com  ← à utiliser dans le frontend",
    title="Valeurs à conserver",
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PARTIE 4 — SERVICE FRONTEND
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "5. Partie 4 — Déploiement du service Frontend", 1)

add_heading(doc, "5.1  Création du service Streamlit Frontend", 2)

add_step(doc, 17, "Créer un second Web Service dans Render")
p = doc.add_paragraph(style="List Number")
p.add_run(
    "Cliquer sur « + New » → « Web Service » → « Deploy an existing image »."
).font.size = Pt(10)

add_step(doc, 18, "Configurer l'image Docker du frontend")

add_table_kv(
    doc,
    [
        ("Image URL", "ghcr.io/rahiming/ml-housting-project/frontend:latest"),
        ("Registry Credential", "ghcr-ml-housing"),
    ],
    col1="Champ",
    col2="Valeur",
)

add_step(doc, 19, "Renseigner les paramètres du service frontend")

add_table_kv(
    doc,
    [
        ("Name", "ml-housing-frontend"),
        ("Region", "Frankfurt (EU Central) — même région que le backend"),
        ("Instance Type", "Free ou Starter"),
        ("Port", "8501"),
    ],
    col1="Champ",
    col2="Valeur",
)

add_step(doc, 20, "Ajouter la variable d'environnement du frontend")
doc.add_paragraph(
    "Le frontend a besoin de connaître l'URL du backend pour envoyer les requêtes "
    "de prédiction."
)

add_table_kv(
    doc,
    [
        (
            "BACKEND_URL",
            "https://ml-housing-backend.onrender.com\n"
            "Remplacer par l'URL réelle notée à l'étape 16",
        ),
    ],
    col1="Variable",
    col2="Valeur",
)

add_step(doc, 21, "Déployer le service frontend")
p = doc.add_paragraph(style="List Number")
p.add_run("Cliquer sur « Deploy Web Service ».").font.size = Pt(10)
p = doc.add_paragraph(style="List Number")
p.add_run("Attendre la fin du déploiement (2 à 5 minutes).").font.size = Pt(10)
p = doc.add_paragraph(style="List Number")
p.add_run(
    "Copier le Service ID (srv-xxxxxxxxxxxxxx) depuis l'onglet « Settings »."
).font.size = Pt(10)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PARTIE 5 — GITHUB CI/CD
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "6. Partie 5 — Configuration GitHub pour le CI/CD", 1)

doc.add_paragraph(
    "Le workflow de déploiement (deploy-production.yml) utilise des secrets et "
    "des variables stockés dans GitHub. Cette section détaille comment les configurer."
)

add_heading(doc, "6.1  Accéder aux paramètres du dépôt", 2)

add_step(doc, 22, "Ouvrir la page de configuration du dépôt")
p = doc.add_paragraph(style="List Number")
p.add_run("Aller sur https://github.com/rahiming/ml-housting-project").font.size = Pt(
    10
)
p = doc.add_paragraph(style="List Number")
p.add_run(
    "Cliquer sur l'onglet « Settings » (en haut de la page du dépôt)."
).font.size = Pt(10)
p = doc.add_paragraph(style="List Number")
p.add_run(
    "Dans le menu de gauche, cliquer sur « Secrets and variables » → « Actions »."
).font.size = Pt(10)

add_heading(doc, "6.2  Créer les secrets GitHub", 2)

doc.add_paragraph(
    "Les secrets sont chiffrés et jamais affichés en clair, même dans les logs CI. "
    "Cliquer sur « New repository secret » pour chaque ligne du tableau suivant."
)

add_table_kv(
    doc,
    [
        ("RENDER_API_KEY", "La clé API Render générée à l'étape 10 (rnd_xxxx…)"),
        ("MINIO_ACCESS_KEY", "L'Access Key ID Cloudflare R2 généré à l'étape 5"),
        ("MINIO_SECRET_KEY", "La Secret Access Key Cloudflare R2 générée à l'étape 5"),
    ],
    col1="Nom du secret",
    col2="Valeur à renseigner",
)

add_heading(doc, "6.3  Créer les variables GitHub", 2)

doc.add_paragraph(
    "Les variables (non-secrètes) sont visibles dans les logs CI. "
    "Cliquer sur l'onglet « Variables » puis « New repository variable » "
    "pour chaque ligne."
)

add_table_kv(
    doc,
    [
        (
            "RENDER_BACKEND_SERVICE_ID",
            "Le Service ID du backend noté à l'étape 16  (srv-xxxxxxxxxxxxxx)",
        ),
        (
            "RENDER_FRONTEND_SERVICE_ID",
            "Le Service ID du frontend noté à l'étape 21  (srv-yyyyyyyyyyyyyy)",
        ),
    ],
    col1="Nom de la variable",
    col2="Valeur à renseigner",
)

add_heading(doc, "6.4  Configurer l'environnement de protection « production »", 2)

add_info_box(
    doc,
    "Le workflow utilise environment: production, ce qui permet d'ajouter une "
    "protection par approbation manuelle avant chaque déploiement. "
    "Cette étape est optionnelle mais recommandée.",
)

add_step(doc, 23, "Créer l'environnement de protection (optionnel)")
p = doc.add_paragraph(style="List Number")
p.add_run(
    "Dans Settings → Environments, cliquer sur « New environment »."
).font.size = Pt(10)
p = doc.add_paragraph(style="List Number")
p.add_run("Name : ").font.size = Pt(10)
p.add_run("production").font.name = "Courier New"
p.runs[-1].font.size = Pt(10)
p = doc.add_paragraph(style="List Number")
p.add_run(
    "Cocher « Required reviewers » et ajouter votre username GitHub. "
    "Tout déploiement nécessitera votre approbation."
).font.size = Pt(10)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PARTIE 6 — PREMIER DÉPLOIEMENT
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "7. Partie 6 — Premier déploiement via le workflow", 1)

add_heading(doc, "7.1  Merger develop → main", 2)

doc.add_paragraph(
    "Le workflow de déploiement ne se déclenche que sur la branche main. "
    "Il faut donc merger la branche develop dans main."
)

add_step(doc, 24, "Créer une Pull Request develop → main")
p = doc.add_paragraph(style="List Number")
p.add_run(
    "Aller sur https://github.com/rahiming/ml-housting-project/compare/main...develop"
).font.size = Pt(10)
p = doc.add_paragraph(style="List Number")
p.add_run("Cliquer sur « Create pull request ».").font.size = Pt(10)
p = doc.add_paragraph(style="List Number")
p.add_run(
    "Attendre que tous les checks CI soient verts (Ruff, Tests, Security…)."
).font.size = Pt(10)
p = doc.add_paragraph(style="List Number")
p.add_run("Cliquer sur « Merge pull request » → « Confirm merge ».").font.size = Pt(10)

add_heading(doc, "7.2  Lancer le déploiement manuellement", 2)

add_step(doc, 25, "Déclencher le workflow Production Deploy")
p = doc.add_paragraph(style="List Number")
p.add_run("Aller sur l'onglet « Actions » du dépôt GitHub.").font.size = Pt(10)
p = doc.add_paragraph(style="List Number")
p.add_run("Dans le menu de gauche, cliquer sur « Production Deploy ».").font.size = Pt(
    10
)
p = doc.add_paragraph(style="List Number")
p.add_run("Cliquer sur « Run workflow » → s'assurer que la branche est ").font.size = (
    Pt(10)
)
p.add_run("main").font.name = "Courier New"
p.runs[-1].font.size = Pt(10)
p.add_run(".").font.size = Pt(10)
p = doc.add_paragraph(style="List Number")
p.add_run("Cliquer sur le bouton vert « Run workflow ».").font.size = Pt(10)

add_heading(doc, "7.3  Suivre l'exécution du workflow", 2)

add_step(doc, 26, "Surveiller les deux jobs du workflow")
doc.add_paragraph("Le workflow exécute deux jobs séquentiels :")

add_table_kv(
    doc,
    [
        (
            "publish-images (5-10 min)",
            "Build Docker des deux images → push vers GHCR avec les tags\n"
            "prod, latest et sha-<commit>",
        ),
        (
            "deploy-render (3-15 min)",
            "Déclenche le déploiement Render via API →\n"
            "Polling du statut toutes les 20s jusqu'à status=live →\n"
            "Affiche les URLs dans le deployment summary",
        ),
    ],
    col1="Job",
    col2="Ce qu'il fait",
)

p = doc.add_paragraph(style="List Number")
p.add_run("Cliquer sur le run en cours pour voir les détails.").font.size = Pt(10)
p = doc.add_paragraph(style="List Number")
p.add_run(
    "À la fin, cliquer sur l'onglet « Summary » pour voir le tableau "
    "récapitulatif avec les URLs de production."
).font.size = Pt(10)

add_success_box(
    doc,
    "Si les deux jobs sont verts, l'application est en production ! "
    "Les URLs backend et frontend sont affichées dans le deployment summary.",
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PARTIE 7 — VÉRIFICATION
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "8. Partie 7 — Vérification et tests", 1)

add_heading(doc, "8.1  Vérification du backend", 2)

add_step(doc, 27, "Tester l'endpoint de santé")
doc.add_paragraph("Ouvrir un terminal et exécuter :")
add_code(
    doc,
    "curl https://ml-housing-backend.onrender.com/health\n\n"
    "# Réponse attendue :\n"
    '{"status": "ok"}',
    title="Test health endpoint",
)
add_warning_box(
    doc,
    "Si le backend est en veille (plan Free), la première requête peut prendre "
    "30 à 60 secondes. Attendre et réessayer si un timeout se produit.",
)

add_step(doc, 28, "Tester l'endpoint de prédiction")
add_code(
    doc,
    "curl -X POST https://ml-housing-backend.onrender.com/predict \\\n"
    '  -H "Content-Type: application/json" \\\n'
    "  -d '{\n"
    '    "median_income": 3.5,\n'
    '    "housing_median_age": 20.0,\n'
    '    "average_rooms": 5.0,\n'
    '    "average_bedrooms": 1.0,\n'
    '    "population": 1000.0,\n'
    '    "average_occupancy": 3.0,\n'
    '    "latitude": 34.0,\n'
    '    "longitude": -118.0\n'
    "  }'\n\n"
    "# Réponse attendue (valeurs approximatives) :\n"
    "{\n"
    '  "prediction": 2.45,\n'
    '  "variant": "A",\n'
    '  "model_version": "model_v1",\n'
    '  "execution_mode": "ab_registry",\n'
    '  "latency_ms": 45.3,\n'
    '  "request_id": "uuid-xxxx-xxxx"\n'
    "}",
    title="Test predict endpoint",
)

add_heading(doc, "8.2  Vérification du frontend", 2)

add_step(doc, 29, "Tester l'interface Streamlit")
p = doc.add_paragraph(style="List Number")
p.add_run(
    "Ouvrir l'URL du frontend dans un navigateur : "
    "https://ml-housing-frontend.onrender.com"
).font.size = Pt(10)
p = doc.add_paragraph(style="List Number")
p.add_run(
    "L'interface doit afficher le titre « A/B Testing - Prediction immobiliere »."
).font.size = Pt(10)
p = doc.add_paragraph(style="List Number")
p.add_run(
    "Renseigner des valeurs dans les champs et cliquer sur « Predire »."
).font.size = Pt(10)
p = doc.add_paragraph(style="List Number")
p.add_run(
    "Vérifier que le résultat s'affiche avec la prédiction, "
    "la variante (A ou B) et la version du modèle."
).font.size = Pt(10)

add_success_box(
    doc,
    "L'application est opérationnelle en production. "
    "Le système A/B routing fonctionne : les utilisateurs avec user_id 'alice' "
    "seront systématiquement routés vers la variante B.",
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PARTIE 8 — DÉPANNAGE
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "9. Partie 8 — Dépannage", 1)

problems = [
    (
        "Le workflow échoue sur « Trigger Backend Deploy » avec une erreur 401",
        "Le secret RENDER_API_KEY est incorrect ou expiré.",
        [
            "Vérifier la valeur dans GitHub Settings → Secrets → RENDER_API_KEY.",
            "Régénérer une clé API dans Render (Account Settings → API Keys).",
            "Mettre à jour le secret GitHub avec la nouvelle valeur.",
        ],
    ),
    (
        "Le workflow échoue sur « Trigger Backend Deploy » avec une erreur 404",
        "Le RENDER_BACKEND_SERVICE_ID est incorrect.",
        [
            "Dans Render, aller sur le service backend → Settings.",
            "Vérifier que le Service ID dans GitHub Variables correspond exactement.",
            "Format attendu : srv-xxxxxxxxxxxxxx (17 caractères après 'srv-').",
        ],
    ),
    (
        "Deploy status reste « build_in_progress » puis « build_failed »",
        "Render ne parvient pas à puller l'image depuis GHCR.",
        [
            "Vérifier que le Registry Credential ghcr-ml-housing est actif dans Render.",
            "Vérifier que le Personal Access Token GitHub n'a pas expiré.",
            "S'assurer que l'image existe bien dans GHCR (onglet Packages du dépôt GitHub).",
            "Dans Render, aller sur le service → Events pour voir les détails de l'erreur.",
        ],
    ),
    (
        "Le backend démarre mais échoue avec « L'application ne peut pas démarrer sans modèle »",
        "Le backend ne peut pas télécharger le modèle depuis Cloudflare R2.",
        [
            "Vérifier que MINIO_ENDPOINT est correct (format https://<id>.r2.cloudflarestorage.com).",
            "Vérifier que MINIO_ACCESS_KEY et MINIO_SECRET_KEY sont bien configurés.",
            "Vérifier que les fichiers model_latest.joblib, model_v1.joblib, model_v2.joblib "
            "sont présents dans le bucket ml-models de Cloudflare R2.",
            "Tester l'accès R2 avec : aws s3 ls s3://ml-models/ "
            "--endpoint-url https://<id>.r2.cloudflarestorage.com",
        ],
    ),
    (
        "Le frontend affiche « Erreur : impossible de contacter le serveur backend FastAPI »",
        "La variable BACKEND_URL est incorrecte ou le backend est en veille.",
        [
            "Vérifier que BACKEND_URL dans le service frontend Render contient l'URL exacte du backend.",
            "Tester manuellement : curl https://ml-housing-backend.onrender.com/health",
            "Si le backend est en veille (plan Free), attendre 60s et réessayer.",
        ],
    ),
    (
        "L'endpoint /predict retourne « legacy_fallback » au lieu de « ab_registry »",
        "Les modèles A/B (model_v1.joblib, model_v2.joblib) ne sont pas chargés.",
        [
            "Vérifier dans les logs Render que les deux fichiers sont téléchargés au démarrage.",
            "Vérifier que MODEL_A_OBJECT_NAME=model_v1.joblib et MODEL_B_OBJECT_NAME=model_v2.joblib "
            "sont définis dans les variables d'environnement Render du backend.",
            "Ce comportement est normal si seul model_latest.joblib est dans R2 "
            "(fallback sur le modèle de référence).",
        ],
    ),
]

for problem, cause, solutions in problems:
    add_heading(doc, problem, 3)
    p = doc.add_paragraph()
    r = p.add_run("Cause probable : ")
    r.bold = True
    r.font.size = Pt(10)
    p.add_run(cause).font.size = Pt(10)
    doc.add_paragraph("Solution :").runs[0].bold = True
    for sol in solutions:
        p2 = doc.add_paragraph(style="List Number")
        p2.add_run(sol).font.size = Pt(10)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# ANNEXES
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "10. Annexes — Récapitulatif complet des variables", 1)

add_heading(doc, "10.1  Variables d'environnement Render — Service Backend", 2)

add_table_kv(
    doc,
    [
        ("MINIO_ENDPOINT", "https://<ACCOUNT_ID>.r2.cloudflarestorage.com"),
        ("MINIO_BUCKET_MODELS", "ml-models"),
        ("MODEL_OBJECT_NAME", "model_latest.joblib"),
        ("MODEL_A_OBJECT_NAME", "model_v1.joblib"),
        ("MODEL_B_OBJECT_NAME", "model_v2.joblib"),
        ("AB_TRAFFIC_B_PERCENT", "50"),
        ("MINIO_ACCESS_KEY", "*** Secret — Access Key ID Cloudflare R2 ***"),
        ("MINIO_SECRET_KEY", "*** Secret — Secret Access Key Cloudflare R2 ***"),
    ],
    col1="Variable",
    col2="Valeur",
)

add_heading(doc, "10.2  Variables d'environnement Render — Service Frontend", 2)

add_table_kv(
    doc,
    [
        ("BACKEND_URL", "https://ml-housing-backend.onrender.com"),
    ],
    col1="Variable",
    col2="Valeur",
)

add_heading(doc, "10.3  Secrets GitHub (Settings → Secrets and variables → Secrets)", 2)

add_table_kv(
    doc,
    [
        ("RENDER_API_KEY", "Clé API Render (rnd_xxxx…)"),
        ("MINIO_ACCESS_KEY", "Access Key ID Cloudflare R2"),
        ("MINIO_SECRET_KEY", "Secret Access Key Cloudflare R2"),
    ],
    col1="Secret",
    col2="Description",
)

add_heading(
    doc, "10.4  Variables GitHub (Settings → Secrets and variables → Variables)", 2
)

add_table_kv(
    doc,
    [
        (
            "RENDER_BACKEND_SERVICE_ID",
            "srv-xxxxxxxxxxxxxx (onglet Settings du service backend Render)",
        ),
        (
            "RENDER_FRONTEND_SERVICE_ID",
            "srv-yyyyyyyyyyyyyy (onglet Settings du service frontend Render)",
        ),
    ],
    col1="Variable",
    col2="Valeur",
)

add_heading(doc, "10.5  URLs utiles", 2)

add_table_kv(
    doc,
    [
        ("Cloudflare Dashboard", "https://dash.cloudflare.com"),
        ("Cloudflare R2 API Tokens", "https://dash.cloudflare.com/profile/api-tokens"),
        ("Render Dashboard", "https://dashboard.render.com"),
        (
            "Render Registry Credentials",
            "https://dashboard.render.com/u/settings#registries",
        ),
        ("Render API Keys", "https://dashboard.render.com/u/settings#apikeys"),
        (
            "GitHub Secrets",
            "https://github.com/rahiming/ml-housting-project/settings/secrets/actions",
        ),
        ("GitHub Actions", "https://github.com/rahiming/ml-housting-project/actions"),
        ("GHCR Images", "https://github.com/rahiming?tab=packages"),
    ],
    col1="Ressource",
    col2="URL",
)

add_heading(doc, "10.6  Ordre récapitulatif des étapes", 2)

all_steps = [
    "Créer un compte Cloudflare",
    "Activer Cloudflare R2 et créer le bucket ml-models",
    "Générer les tokens d'accès R2 et noter Access Key ID + Secret",
    "Uploader model_latest.joblib, model_v1.joblib, model_v2.joblib dans R2",
    "Créer un compte Render (via GitHub OAuth)",
    "Créer un PAT GitHub (scope read:packages) et un Registry Credential dans Render",
    "Générer une clé API Render",
    "Créer le service backend dans Render (image GHCR + variables d'env)",
    "Créer le service frontend dans Render (image GHCR + BACKEND_URL)",
    "Copier les deux Service IDs Render",
    "Créer les secrets GitHub : RENDER_API_KEY, MINIO_ACCESS_KEY, MINIO_SECRET_KEY",
    "Créer les variables GitHub : RENDER_BACKEND_SERVICE_ID, RENDER_FRONTEND_SERVICE_ID",
    "Merger develop → main via Pull Request",
    "Lancer manuellement le workflow « Production Deploy » depuis l'onglet Actions",
    "Vérifier que les deux jobs sont verts et tester les URLs de production",
]

table = doc.add_table(rows=len(all_steps), cols=2)
table.style = "Table Grid"
for i, step_text in enumerate(all_steps, 1):
    cells = table.rows[i - 1].cells
    cells[0].text = str(i)
    cells[1].text = step_text
    shading = OxmlElement("w:shd")
    shading.set(
        qn("w:fill"),
        (
            "1A56DB"
            if i <= 4
            else ("059669" if i <= 7 else ("7C3AED" if i <= 12 else "DC2626"))
        ),
    )
    shading.set(qn("w:val"), "clear")
    cells[0]._tc.get_or_add_tcPr().append(shading)
    for run in cells[0].paragraphs[0].runs:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.bold = True
        run.font.size = Pt(10)
    for run in cells[1].paragraphs[0].runs:
        run.font.size = Pt(9)
    if i % 2 == 0:
        shading2 = OxmlElement("w:shd")
        shading2.set(qn("w:fill"), "F9FAFB")
        shading2.set(qn("w:val"), "clear")
        cells[1]._tc.get_or_add_tcPr().append(shading2)

# ── Pied de page ──────────────────────────────────────────────────────────────

doc.add_paragraph()
footer_para = doc.add_paragraph()
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer_para.add_run(
    f"Document généré le {datetime.date.today().strftime('%d/%m/%Y')}  •  "
    "Projet ML Housing Prediction  •  Pour usage interne"
)
r.font.size = Pt(8)
r.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
r.italic = True

# ── Sauvegarde ────────────────────────────────────────────────────────────────

output_path = os.path.join(
    os.path.dirname(__file__), "..", "docs", "Guide_Deploiement_Render.docx"
)
output_path = os.path.normpath(output_path)
doc.save(output_path)
print(f"Document généré : {output_path}")
