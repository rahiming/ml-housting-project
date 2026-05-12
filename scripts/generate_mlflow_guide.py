"""Génère le guide MLFlow au format .docx."""

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

doc = Document()


# ── Helpers ───────────────────────────────────────────────────────────────────
def set_font(run, bold=False, italic=False, size=11, color=None, name="Calibri"):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = name
    if color:
        run.font.color.rgb = RGBColor(*color)


def h(level, text):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(14 if level <= 2 else 8)
    p.paragraph_format.space_after = Pt(4)
    return p


def para(text="", bold_prefix=None, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        set_font(r, bold=True)
    if text:
        r = p.add_run(text)
        set_font(r)
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_font(r)


def numbered(text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_font(r)


def shade_para(p, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    p._p.get_or_add_pPr().append(shd)


def code(lines):
    for line in lines.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line if line else " ")
        r.font.name = "Courier New"
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
        shade_para(p, "F0F0F0")
    doc.add_paragraph()


def tip(label, text, fill="FFF3CD"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(label + "  ")
    set_font(
        r1,
        bold=True,
        color=(0x85, 0x63, 0x04) if fill == "FFF3CD" else (0x0C, 0x54, 0x6D),
    )
    r2 = p.add_run(text)
    set_font(r2, color=(0x49, 0x33, 0x00) if fill == "FFF3CD" else (0x0C, 0x54, 0x6D))
    shade_para(p, fill)


def important(label, text):
    tip(label, text, fill="F8D7DA")


def schema_line(text, color=(0x1F, 0x49, 0x7D)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(*color)
    shade_para(p, "EBF3FB")


# ═══════════════════════════════════════════════════════════════════════════════
# TITRE
# ═══════════════════════════════════════════════════════════════════════════════
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_before = Pt(60)
r = tp.add_run("Guide MLFlow — Traçabilité & Déploiement")
r.bold = True
r.font.size = Pt(24)
r.font.name = "Calibri"
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sp.add_run("Projet fil rouge — Industrialisation ML")
r.font.size = Pt(14)
r.font.name = "Calibri"
r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

sp2 = doc.add_paragraph()
sp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sp2.add_run(
    "De l'expérimentation locale à la mise en production sur Render via HuggingFace Hub"
)
r.font.size = Pt(12)
r.italic = True
r.font.name = "Calibri"
r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SOMMAIRE
# ═══════════════════════════════════════════════════════════════════════════════
h(1, "Sommaire")
toc = [
    ("Partie 1", "MLOps et le rôle de MLFlow — Vision d'ensemble"),
    ("Partie 2", "Architecture concrète de ce projet"),
    ("Partie 3", "Installation et configuration"),
    ("Partie 4", "Tracking dans le pipeline principal (src/training/pipeline.py)"),
    ("Partie 5", "Tracking dans le notebook A/B (notebooks/02_train_ab_models.ipynb)"),
    ("Partie 6", "L'interface MLFlow UI — Lire et comparer les runs"),
    ("Partie 7", "Model Registry — La porte de mise en production"),
    ("Partie 8", "Connexion Registry → HuggingFace Hub → Render"),
    ("Partie 9", "Récapitulatif des fichiers du projet"),
    ("Annexe", "Aide-mémoire des commandes et flux de soutenance"),
]
for num, title in toc:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(num + "  —  ")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.name = "Calibri"
    r2 = p.add_run(title)
    r2.font.size = Pt(11)
    r2.font.name = "Calibri"

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# PARTIE 1 — MLOPS ET RÔLE DE MLFLOW
# ═══════════════════════════════════════════════════════════════════════════════
h(1, "Partie 1 — MLOps et le rôle de MLFlow")

h(2, "1.1  Qu'est-ce que MLOps ?")
para(
    "Le terme MLOps (Machine Learning Operations) désigne l'ensemble des pratiques qui "
    "permettent de passer d'un modèle développé en notebook à un service disponible en "
    "production, maintenable dans le temps et reproductible par n'importe quel membre de "
    "l'équipe."
)
para(
    "Sans MLOps, un projet ML ressemble à ceci : un data scientist entraîne un modèle sur "
    "son laptop, copie le fichier .joblib quelque part, et le déploiement se fait à la main "
    "sans trace de ce qui a été fait. Si le modèle se dégrade en production deux mois plus "
    "tard, impossible de savoir avec quels paramètres il avait été entraîné, ni de reproduire "
    "les conditions d'origine."
)
para(
    "MLOps répond à cette question fondamentale : comment passer du code qui tourne une fois "
    "sur un ordinateur au service qui tourne en continu, de manière fiable et auditée ?"
)

h(2, "1.2  Le cycle de vie d'un modèle ML")
para("Un modèle ML passe par plusieurs étapes avant d'arriver en production :")

schema_line("┌─────────────────────────────────────────────────────────────────┐")
schema_line("│  Développement  →  Expérimentation  →  Validation  →  Production │")
schema_line("└─────────────────────────────────────────────────────────────────┘")
doc.add_paragraph()

stages_mlops = [
    (
        "Développement",
        "Le data scientist explore les données, choisit un algorithme, teste "
        "des hyperparamètres. C'est la phase notebook.",
    ),
    (
        "Expérimentation",
        "On lance plusieurs entraînements avec des paramètres différents et on compare "
        "les résultats. MLFlow Tracking intervient ici.",
    ),
    (
        "Validation",
        "On sélectionne le meilleur modèle et on le soumet à une validation formelle "
        "(métriques seuil, revue humaine). MLFlow Model Registry intervient ici.",
    ),
    (
        "Production",
        "Le modèle validé est déployé sur le service. Dans ce projet : upload vers "
        "HuggingFace Hub, puis chargé par le backend Render.",
    ),
    (
        "Monitoring",
        "En production, on surveille les prédictions pour détecter une dérive du modèle. "
        "Dans ce projet : les logs PostgreSQL et le notebook 03 d'analyse A/B.",
    ),
]
for stage, desc in stages_mlops:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run("▸  " + stage + " : ")
    set_font(r1, bold=True, color=(0x1F, 0x49, 0x7D))
    r2 = p.add_run(desc)
    set_font(r2)

h(2, "1.3  Pourquoi MLFlow spécifiquement ?")
para(
    "MLFlow est l'outil open-source le plus utilisé dans l'industrie pour la gestion du cycle "
    "de vie ML. Il est utilisé par Netflix, Databricks, Booking.com, et de nombreuses équipes "
    "data. Il couvre trois besoins critiques :"
)
bullet("Traçabilité : retrouver exactement comment un modèle a été produit.")
bullet(
    "Comparaison : identifier le meilleur modèle parmi plusieurs expériences en un coup d'œil."
)
bullet(
    "Gouvernance : s'assurer qu'un modèle ne part pas en production sans validation explicite."
)

h(2, "1.4  Les 4 composants de MLFlow")
concepts = [
    (
        "MLFlow Tracking",
        "Enregistre automatiquement les paramètres, métriques et artefacts de chaque run "
        "d'entraînement. C'est le composant central utilisé dans ce projet.",
    ),
    (
        "MLFlow Projects",
        "Permet de packager un projet ML de manière reproductible (fichier MLproject). "
        "Non utilisé dans ce projet.",
    ),
    (
        "MLFlow Models",
        "Format standardisé pour sauvegarder et charger un modèle quel que soit "
        "le framework (sklearn, pytorch, tensorflow, etc.).",
    ),
    (
        "MLFlow Model Registry",
        "Registre centralisé pour versionner les modèles et les promouvoir de manière "
        "contrôlée : None → Staging → Production → Archived. C'est la porte d'entrée "
        "vers le déploiement dans ce projet.",
    ),
]
for name, desc in concepts:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run("▸  " + name + " : ")
    set_font(r1, bold=True, color=(0x1F, 0x49, 0x7D))
    r2 = p.add_run(desc)
    set_font(r2)

h(2, "1.5  Vocabulaire essentiel")
vocab = [
    ("Experiment", 'Groupe de runs liés à un même objectif. Ex : "housing_ab_models".'),
    (
        "Run",
        "Une exécution unique du pipeline. Chaque appel à mlflow.start_run() crée un run.",
    ),
    (
        "Parameter",
        "Valeur fixée AVANT l'entraînement. Ex : n_estimators=100, learning_rate=0.05.",
    ),
    (
        "Metric",
        "Valeur mesurée APRÈS l'entraînement. Ex : mae=0.41, rmse=0.59, r2=0.74.",
    ),
    (
        "Artifact",
        "Fichier produit par le run. Ex : le modèle sérialisé, le fichier metrics.json.",
    ),
    (
        "Registry",
        "Registre de modèles versionnés. Un modèle du Registry a un nom, une version, "
        "et un stage (None / Staging / Production / Archived).",
    ),
    ("Tag", "Métadonnée libre associée à un run. Ex : auteur, environnement."),
]
for term, definition in vocab:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run(term.ljust(14))
    set_font(r1, bold=True, name="Courier New", size=10)
    r2 = p.add_run(" — " + definition)
    set_font(r2)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# PARTIE 2 — ARCHITECTURE DU PROJET
# ═══════════════════════════════════════════════════════════════════════════════
h(1, "Partie 2 — Architecture concrète de ce projet")

h(2, "2.1  Vue d'ensemble du pipeline complet")
para(
    "Ce projet implémente un pipeline MLOps complet. MLFlow n'est pas un outil isolé : "
    "il est le centre de décision qui conditionne ce qui part ou non en production."
)

schema_line("  DÉVELOPPEMENT LOCAL")
schema_line("  ┌──────────────────────────────────────────────────────────────┐")
schema_line("  │  notebooks/02_train_ab_models.ipynb                          │")
schema_line("  │       ↓ mlflow.log_params / log_metrics / log_model          │")
schema_line("  │  src/training/pipeline.py (python main.py)                   │")
schema_line("  │       ↓ mlflow.log_param / log_metric / log_model            │")
schema_line("  │  MLFlow Tracking (mlruns/)  ←  tous les runs loggés ici      │")
schema_line("  │       ↓                                                       │")
schema_line("  │  MLFlow Model Registry  ←  promotion en stage Production      │")
schema_line("  └──────────────────────────────────────────────────────────────┘")
schema_line("       ↓  seulement si stage == Production")
schema_line("  ┌──────────────────────────────────────────────────────────────┐")
schema_line("  │  scripts/upload_model_to_hf.py                               │")
schema_line("  │       Charge housing_model_A/Production → model_v1.joblib    │")
schema_line("  │       Charge housing_model_B/Production → model_v2.joblib    │")
schema_line("  │       Charge meilleur R²               → model_latest.joblib │")
schema_line("  │       Upload vers HuggingFace Hub                            │")
schema_line("  └──────────────────────────────────────────────────────────────┘")
schema_line("       ↓")
schema_line("  ┌──────────────────────────────────────────────────────────────┐")
schema_line("  │  Backend FastAPI sur Render                                  │")
schema_line("  │       Télécharge les .joblib au démarrage depuis HF Hub      │")
schema_line("  │       Sert les prédictions A/B en production                 │")
schema_line("  └──────────────────────────────────────────────────────────────┘")
doc.add_paragraph()

h(2, "2.2  Ce que MLFlow apporte concrètement dans ce projet")
before_after = [
    (
        "Comparaison des runs",
        "Ouvrir chaque fichier JSON manuellement dans artifacts/metrics/",
        "Interface MLFlow UI avec tableau comparatif automatique et graphiques",
    ),
    (
        "Identifier le meilleur modèle",
        "Lire les JSON et comparer à la main",
        'client.search_runs(order_by=["metrics.r2 DESC"]) → résultat immédiat',
    ),
    (
        "Gate de production",
        "Tout .joblib dans artifacts/ pouvait être uploadé sans validation",
        "Seul un modèle en stage Production dans le Registry est uploadé sur HF Hub",
    ),
    (
        "Reproductibilité",
        "Aucune trace des hyperparamètres utilisés",
        "Chaque run enregistre ses paramètres, métriques, et le modèle associé",
    ),
    (
        "Versioning",
        "Fichiers model_v1.joblib, model_v2.joblib renommés manuellement",
        "MLFlow versionne automatiquement chaque registered model",
    ),
]

table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "Problème"
hdr[1].text = "Avant MLFlow"
hdr[2].text = "Après MLFlow"
for cell in hdr:
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)
for aspect, before, after in before_after:
    rc = table.add_row().cells
    rc[0].text = aspect
    rc[1].text = before
    rc[2].text = after
    for c in rc:
        c.paragraphs[0].runs[0].font.size = Pt(9)
doc.add_paragraph()

h(2, "2.3  Ce que MLFlow ne remplace PAS dans ce projet")
para("Il est important de comprendre les limites de MLFlow dans cette architecture :")
bullet(
    "MLFlow ne remplace pas les fichiers .joblib dans artifacts/. "
    "Le notebook 02 continue de sauvegarder model_v1.joblib et model_v2.joblib "
    "pour la compatibilité avec les autres notebooks. MLFlow s'ajoute, il ne supprime pas."
)
bullet(
    "MLFlow n'est pas installé dans le backend Render. Le backend charge "
    "les .joblib directement depuis HuggingFace Hub. MLFlow est un outil de développement."
)
bullet(
    "MLFlow ne monitore pas la production. La surveillance en production est assurée "
    "par les logs PostgreSQL (Neon) et le notebook 03 d'analyse A/B."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# PARTIE 3 — INSTALLATION
# ═══════════════════════════════════════════════════════════════════════════════
h(1, "Partie 3 — Installation et configuration")

h(2, "3.1  Installer MLFlow")
para("Dans le terminal, avec l'environnement virtuel activé :")
code("pip install mlflow==2.22.0")
para("Ajouter la dépendance dans requirements.txt (racine du projet, PAS backend/) :")
code("mlflow==2.22.0")
tip(
    "ℹ️ Note :",
    "MLFlow n'est PAS ajouté à backend/requirements.txt. "
    "Le backend Render n'a besoin que de charger un .joblib — il ne fait pas de tracking. "
    "Ajouter MLFlow au backend alourdirait l'image Docker inutilement.",
)

h(2, "3.2  Le Tracking URI — concept clé")
para(
    "Le Tracking URI indique à MLFlow où stocker les données de runs. "
    "Par défaut, MLFlow utilise le répertoire courant au moment où le script s'exécute. "
    "C'est un point critique quand on travaille avec des notebooks Jupyter."
)
para(
    "Problème concret : si on lance un notebook depuis notebooks/, MLFlow crée "
    "notebooks/mlruns/ au lieu du dossier mlruns/ à la racine du projet. "
    "On se retrouve alors avec deux bases de données séparées."
)
para("Solution : toujours spécifier explicitement le Tracking URI dans les notebooks :")
code(
    "# Dans les notebooks — pointer vers le mlruns/ à la racine du projet\n"
    "mlflow.set_tracking_uri((Path('..') / 'mlruns').resolve().as_uri())"
)
para(
    "Cette ligne doit apparaître AVANT tout appel à set_experiment() ou start_run(). "
    "Pour les scripts Python (pipeline.py, upload_model_to_hf.py), le répertoire courant "
    "est toujours la racine du projet, donc le Tracking URI par défaut est correct."
)
tip(
    "⚠️ Piège fréquent :",
    "Oublier cette ligne dans un notebook fait que les runs sont enregistrés dans "
    "notebooks/mlruns/ au lieu de mlruns/. La commande 'mlflow ui' lancée depuis la racine "
    "ne les verra pas. Vérifier toujours que tous les runs apparaissent dans la même UI.",
)

h(2, "3.3  Mettre à jour .gitignore")
para("Vérifier que .gitignore contient ces lignes (déjà en place dans ce projet) :")
code("# MLFlow\nmlruns/\nmlartifacts/")
para(
    "Le dossier mlruns/ contient la base de données locale de MLFlow. "
    "Il peut devenir très volumineux (modèles sérialisés, artefacts) "
    "et ne doit jamais être commité dans Git."
)

h(2, "3.4  Vérifier l'installation")
code(
    'python -c "import mlflow; print(mlflow.__version__)"\n# Résultat attendu : 2.22.0'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# PARTIE 4 — PIPELINE PRINCIPAL
# ═══════════════════════════════════════════════════════════════════════════════
h(1, "Partie 4 — Tracking dans le pipeline principal")

para(
    "Le fichier src/training/pipeline.py orchestre chargement → entraînement → évaluation "
    "→ sauvegarde. C'est là qu'est entraîné le modèle de production (RandomForest sur "
    "les données California Housing). MLFlow y est intégré pour tracer chaque run "
    "de ce pipeline."
)

h(2, "4.1  Pourquoi tracer ce pipeline ?")
para(
    "Ce pipeline est relancé à chaque fois qu'on veut entraîner une nouvelle version "
    "du modèle de production. Sans MLFlow, il faut ouvrir manuellement les fichiers JSON "
    "pour comparer les métriques de v1, v2, v3... Avec MLFlow, toutes les versions sont "
    "comparables en un tableau dans l'interface."
)
para(
    "Autre bénéfice : si un modèle déployé en production se met à donner de mauvaises "
    "prédictions, MLFlow permet de retrouver exactement les paramètres du modèle précédent "
    "et de le réentraîner à l'identique."
)

h(2, "4.2  Structure de la fonction run_pipeline() avec MLFlow")
para("Voici le code intégral de src/training/pipeline.py (version finale du projet) :")
code(
    '"""Orchestration du pipeline ML avec tracking MLFlow."""\n'
    "\n"
    "import json\n"
    "import logging\n"
    "from pathlib import Path\n"
    "\n"
    "import joblib\n"
    "import mlflow\n"
    "import mlflow.sklearn\n"
    "\n"
    "from src.common.features import split_features_target, split_train_test\n"
    "from src.training.data import load_housing_data\n"
    "from src.training.evaluate import evaluate_model\n"
    "from src.training.train import train_model\n"
    "\n"
    "logger = logging.getLogger(__name__)\n"
    'MLFLOW_EXPERIMENT = "housing_price_prediction"\n'
    "\n"
    "\n"
    'def run_pipeline(artifacts_dir: str = "artifacts", version: str = "") -> dict:\n'
    "    ...\n"
    "    # Charger les données et splitter\n"
    "    df = load_housing_data()\n"
    "    X, y = split_features_target(df)\n"
    "    X_train, X_test, y_train, y_test = split_train_test(X, y)\n"
    "\n"
    "    # Sélectionner ou créer l'expérience MLFlow\n"
    "    mlflow.set_experiment(MLFLOW_EXPERIMENT)\n"
    "\n"
    '    with mlflow.start_run(run_name=f"run_{version}" if version else None):\n'
    "\n"
    "        # 1. Enregistrer les paramètres AVANT l'entraînement\n"
    '        mlflow.log_param("model_type",   "RandomForestRegressor")\n'
    '        mlflow.log_param("n_estimators", 100)\n'
    '        mlflow.log_param("random_state", 42)\n'
    '        mlflow.log_param("version",      version or "latest")\n'
    '        mlflow.log_param("train_size",   len(X_train))\n'
    '        mlflow.log_param("test_size",    len(X_test))\n'
    "\n"
    "        # 2. Entraîner\n"
    "        model = train_model(X_train, y_train)\n"
    "\n"
    "        # 3. Évaluer et enregistrer les métriques APRÈS l'entraînement\n"
    "        metrics = evaluate_model(model, X_test, y_test)\n"
    '        mlflow.log_metric("mae",  metrics["mae"])\n'
    '        mlflow.log_metric("rmse", metrics["rmse"])\n'
    '        mlflow.log_metric("r2",   metrics["r2"])\n'
    "\n"
    "        # 4. Sauvegarder le .joblib (chemin existant conservé)\n"
    '        suffix = f"_{version}" if version else ""\n'
    '        model_file = models_path / f"model{suffix}.joblib"\n'
    "        joblib.dump(model, model_file)\n"
    '        joblib.dump(model, models_path / "model_latest.joblib")\n'
    "\n"
    "        # 5. Enregistrer le modèle dans MLFlow et dans le Registry\n"
    "        mlflow.sklearn.log_model(\n"
    "            sk_model=model,\n"
    '            artifact_path="model",\n'
    "            registered_model_name=(\n"
    '                f"housing_{version}" if version else "housing_latest"\n'
    "            ),\n"
    "        )\n"
    "        # 6. Joindre le fichier de métriques JSON comme artefact\n"
    '        mlflow.log_artifact(str(metrics_file), artifact_path="metrics")\n'
    "\n"
    "    return metrics"
)

h(2, "4.3  Explication de chaque appel MLFlow")
explanations = [
    (
        "mlflow.set_experiment(MLFLOW_EXPERIMENT)",
        "Crée l'expérience 'housing_price_prediction' si elle n'existe pas, ou la "
        "sélectionne si elle existe déjà. Tous les runs du pipeline principal seront "
        "regroupés sous cette expérience dans l'UI. Cela permet de les isoler des "
        "runs du notebook A/B.",
    ),
    (
        "with mlflow.start_run(run_name=...):",
        "Ouvre un contexte de run MLFlow. Tout ce qui est logué à l'intérieur du bloc "
        "with est associé à ce run unique. À la sortie du with, le run est automatiquement "
        "clôturé et sauvegardé dans mlruns/. Si le run_name est fourni (ex: 'run_v3'), "
        "il apparaîtra avec ce nom dans l'interface.",
    ),
    (
        "mlflow.log_param(clé, valeur)",
        "Enregistre un paramètre : une valeur fixée AVANT l'entraînement. "
        "Ici on logue les hyperparamètres du RandomForest (n_estimators, random_state) "
        "et des informations contextuelles (train_size, test_size, version). "
        "Ces paramètres permettent de comprendre pourquoi les métriques diffèrent entre runs.",
    ),
    (
        "mlflow.log_metric(clé, valeur)",
        "Enregistre une métrique : un résultat numérique mesuré APRÈS l'entraînement. "
        "Ici mae (erreur absolue moyenne), rmse (racine de l'erreur quadratique), "
        "et r2 (coefficient de détermination). C'est sur ces métriques qu'on base "
        "la décision de promouvoir ou non un modèle.",
    ),
    (
        "mlflow.sklearn.log_model(..., registered_model_name=...)",
        "Sauvegarde le pipeline sklearn complet (preprocessing + modèle) dans les "
        "artefacts MLFlow ET l'enregistre dans le Model Registry sous le nom donné. "
        "C'est ce qui permet de faire mlflow.sklearn.load_model('models:/housing_latest/Production') "
        "plus tard pour charger le modèle directement depuis le Registry.",
    ),
    (
        "mlflow.log_artifact(chemin_fichier)",
        "Copie un fichier existant (ici metrics.json) dans le stockage MLFlow. "
        "Ce fichier devient visible et téléchargeable depuis l'UI dans la section Artifacts du run. "
        "Utile pour joindre des fichiers de configuration, des rapports, des plots.",
    ),
]
for fn, expl in explanations:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run(fn + "\n")
    set_font(r1, bold=True, name="Courier New", size=9, color=(0x1F, 0x49, 0x7D))
    r2 = p.add_run("    → " + expl)
    set_font(r2, color=(0x40, 0x40, 0x40))

h(2, "4.4  Tester")
para("Depuis la racine du projet, lancer le pipeline :")
code("python main.py")
para("Un dossier mlruns/ est créé à la racine. Résultat attendu dans la console :")
code(
    "INFO | pipeline | Modele entraine avec succes.\n"
    "INFO | pipeline | Evaluation terminee. metrics={'mae': 0.33, 'rmse': 0.50, 'r2': 0.81}\n"
    "INFO | pipeline | Run MLFlow enregistre avec succes."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# PARTIE 5 — NOTEBOOK A/B
# ═══════════════════════════════════════════════════════════════════════════════
h(1, "Partie 5 — Tracking dans le notebook A/B")

para(
    "Le notebook notebooks/02_train_ab_models.ipynb entraîne deux modèles en parallèle : "
    "le modèle A (RandomForestRegressor, rôle champion) et le modèle B "
    "(GradientBoostingRegressor, rôle challenger). MLFlow permet ici de comparer "
    "directement les deux architectures dans la même interface."
)

h(2, "5.1  Cellule 1 — Imports")
para("Ajouter mlflow aux imports en respectant l'ordre alphabétique (standard PEP8) :")
code(
    "# Cellule 1 - Charger les bibliothèques\n"
    "from pathlib import Path\n"
    "import json\n"
    "import shutil\n"
    "import joblib\n"
    "import mlflow\n"
    "import mlflow.sklearn\n"
    "import numpy as np\n"
    "\n"
    "from sklearn.datasets import fetch_california_housing\n"
    "from sklearn.ensemble import GradientBoostingRegressor, RandomForestRegressor\n"
    "from sklearn.metrics import mean_absolute_error, mean_squared_error, r2_score\n"
    "from sklearn.model_selection import train_test_split\n"
    "from sklearn.pipeline import Pipeline\n"
    "from sklearn.preprocessing import StandardScaler"
)

h(2, "5.2  Cellule 2b — Configurer MLFlow (ligne critique)")
para(
    "Insérer cette cellule après le chargement du dataset et AVANT les cellules "
    "d'entraînement. Deux lignes sont nécessaires, et leur ordre est important."
)
code(
    "# Cellule 2b — Configurer MLFlow\n"
    "\n"
    "# LIGNE CRITIQUE : pointer vers le mlruns/ à la racine du projet\n"
    "# Sans cette ligne, Jupyter crée un notebooks/mlruns/ séparé\n"
    "# et les runs n'apparaissent pas dans la même UI que le pipeline\n"
    "mlflow.set_tracking_uri((Path('..') / 'mlruns').resolve().as_uri())\n"
    "\n"
    "# Créer ou sélectionner l'expérience dédiée aux modèles A/B\n"
    'mlflow.set_experiment("housing_ab_models")\n'
    'print("Expérience MLFlow configurée : housing_ab_models")'
)
tip(
    "⚠️ Pourquoi set_tracking_uri est indispensable ici :",
    "Jupyter exécute les cellules depuis le dossier notebooks/. "
    "Sans set_tracking_uri(), MLFlow crée un notebooks/mlruns/ au lieu du mlruns/ "
    "à la racine. Résultat : les runs du notebook ne sont pas visibles dans l'UI, "
    "et upload_model_to_hf.py ne trouve pas les modèles enregistrés. "
    "La méthode resolve().as_uri() garantit un chemin absolu quel que soit l'OS.",
)

h(2, "5.3  Cellule 4 — Entraînement du modèle A (champion)")
para(
    "Envelopper l'entraînement dans un contexte MLFlow. On utilise log_params() "
    "(avec un 's') pour passer un dictionnaire complet en une seule ligne."
)
code(
    "# Cellule 4 - Entraîner le modèle A (champion) avec MLFlow\n"
    'with mlflow.start_run(run_name="model_A_RandomForest"):\n'
    "\n"
    "    params_a = {\n"
    '        "model_type":   "RandomForestRegressor",\n'
    '        "n_estimators": 30,\n'
    '        "max_depth":    8,\n'
    '        "random_state": 42,\n'
    '        "role":         "champion",\n'
    "    }\n"
    "    mlflow.log_params(params_a)  # logue tout le dict d'un coup\n"
    "\n"
    "    model_a = Pipeline([\n"
    '        ("scaler", StandardScaler()),\n'
    '        ("model",  RandomForestRegressor(\n'
    '            n_estimators=params_a["n_estimators"],\n'
    '            max_depth=params_a["max_depth"],\n'
    '            random_state=params_a["random_state"],\n'
    "            n_jobs=-1,\n"
    "        )),\n"
    "    ])\n"
    "    model_a.fit(X_train, y_train)\n"
    "\n"
    "    metrics_a = evaluate_pipeline(model_a, X_test, y_test)\n"
    "    mlflow.log_metrics(metrics_a)  # logue mae, rmse, r2 d'un coup\n"
    "\n"
    "    mlflow.sklearn.log_model(\n"
    "        model_a,\n"
    '        artifact_path="model",\n'
    "        # registered_model_name enregistre dans le Model Registry\n"
    '        registered_model_name="housing_model_A",\n'
    "    )\n"
    '    print("Modèle A :", metrics_a)'
)

h(2, "5.4  Cellule 5 — Entraînement du modèle B (challenger)")
para(
    "Même structure que le modèle A. Le registered_model_name différent permet "
    "de distinguer les deux modèles dans le Registry."
)
code(
    "# Cellule 5 - Entraîner le modèle B (challenger) avec MLFlow\n"
    'with mlflow.start_run(run_name="model_B_GradientBoosting"):\n'
    "\n"
    "    params_b = {\n"
    '        "model_type":    "GradientBoostingRegressor",\n'
    '        "n_estimators":  80,\n'
    '        "learning_rate": 0.05,\n'
    '        "random_state":  42,\n'
    '        "role":          "challenger",\n'
    "    }\n"
    "    mlflow.log_params(params_b)\n"
    "\n"
    "    model_b = Pipeline([\n"
    '        ("scaler", StandardScaler()),\n'
    '        ("model",  GradientBoostingRegressor(\n'
    '            n_estimators=params_b["n_estimators"],\n'
    '            learning_rate=params_b["learning_rate"],\n'
    '            random_state=params_b["random_state"],\n'
    "        )),\n"
    "    ])\n"
    "    model_b.fit(X_train, y_train)\n"
    "\n"
    "    metrics_b = evaluate_pipeline(model_b, X_test, y_test)\n"
    "    mlflow.log_metrics(metrics_b)\n"
    "\n"
    "    mlflow.sklearn.log_model(\n"
    "        model_b,\n"
    '        artifact_path="model",\n'
    '        registered_model_name="housing_model_B",\n'
    "    )\n"
    '    print("Modèle B :", metrics_b)'
)

h(2, "5.5  Cellule 6 — Sauvegarde .joblib (inchangée)")
para(
    "La cellule qui sauvegarde model_v1.joblib et model_v2.joblib dans artifacts/models/ "
    "n'est pas modifiée. MLFlow conserve ses propres copies sérialisées dans mlruns/. "
    "Les deux coexistent sans conflit et servent des usages différents :"
)
bullet(
    "artifacts/models/*.joblib : utilisés par les notebooks locaux (03_ab_analysis) "
    "et potentiellement par d'autres scripts."
)
bullet(
    "mlruns/ : utilisé par MLFlow UI pour l'exploration, et par le Registry "
    "pour la promotion vers la production."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# PARTIE 6 — UI
# ═══════════════════════════════════════════════════════════════════════════════
h(1, "Partie 6 — L'interface MLFlow UI")

h(2, "6.1  Lancer l'interface")
para("Dans un terminal, depuis la racine du projet :")
code("mlflow ui")
para("Résultat dans le terminal :")
code("[INFO] Starting gunicorn 21.2.0\n[INFO] Listening at: http://127.0.0.1:5000")
para("Ouvrir un navigateur à l'adresse : http://127.0.0.1:5000")
tip(
    "⚠️ Important :",
    "Ce serveur est local uniquement. Il doit être lancé depuis la racine du projet "
    "(là où se trouve le dossier mlruns/). Si vous le lancez depuis notebooks/, "
    "il lira notebooks/mlruns/ et ne verra pas les runs du pipeline principal.",
)
tip(
    "ℹ️ macOS / Windows :",
    "Si le port 5000 est occupé, utiliser : mlflow ui --port 5001",
    fill="D1ECF1",
)

h(2, "6.2  Naviguer dans l'interface")
bullet(
    "Colonne gauche : liste des expériences. "
    "Vous devez voir : housing_price_prediction et housing_ab_models."
)
bullet(
    "Zone centrale : tableau de tous les runs avec leurs paramètres et métriques. "
    "Chaque ligne est un run (= une exécution du pipeline ou du notebook)."
)
bullet("Bouton « Compare » : graphique de comparaison entre runs sélectionnés.")
bullet(
    "Onglet « Models » (en haut) : registre de tous les modèles versionnés. "
    "Vous devez voir : housing_model_A, housing_model_B, et housing_latest."
)

h(2, "6.3  Comparer les modèles A et B")
numbered(
    "Dans l'expérience housing_ab_models, cocher les cases de model_A_RandomForest "
    "et model_B_GradientBoosting."
)
numbered("Cliquer sur le bouton bleu « Compare ».")
numbered(
    "MLFlow affiche un tableau côte-à-côte des paramètres et métriques des deux runs."
)
numbered(
    "Analyser : R² plus élevé = meilleur pouvoir explicatif. MAE plus faible = prédictions "
    "plus précises. Sur ce dataset, model_A (RandomForest) obtient généralement un R² ~0.74 "
    "et model_B (GradientBoosting) ~0.72."
)

h(2, "6.4  Identifier le meilleur run depuis le code Python")
para(
    "Pour automatiser la sélection sans passer par l'interface, on peut interroger "
    "MLFlow programmatiquement. C'est ce que fait le script upload_model_to_hf.py "
    "pour choisir le modèle latest :"
)
code(
    "from mlflow.tracking import MlflowClient\n"
    "\n"
    "client = MlflowClient()\n"
    "\n"
    "# Récupérer l'ID de l'expérience\n"
    'experiment = client.get_experiment_by_name("housing_price_prediction")\n'
    "\n"
    "# Chercher le meilleur run selon r2 décroissant\n"
    "best_run = client.search_runs(\n"
    "    experiment_ids=[experiment.experiment_id],\n"
    '    filter_string="metrics.r2 > 0",\n'
    '    order_by=["metrics.r2 DESC"],\n'
    "    max_results=1,\n"
    ")[0]\n"
    "\n"
    'print("Run ID  :", best_run.info.run_id)\n'
    'print("R²      :", best_run.data.metrics["r2"])\n'
    'print("MAE     :", best_run.data.metrics["mae"])\n'
    'print("Params  :", best_run.data.params)'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# PARTIE 7 — MODEL REGISTRY
# ═══════════════════════════════════════════════════════════════════════════════
h(1, "Partie 7 — MLFlow Model Registry : la porte de production")

para(
    "Le Model Registry est le composant le plus important dans une perspective MLOps. "
    "C'est lui qui sépare le monde du développement du monde de la production. "
    "Un modèle ne peut partir en production que s'il a été explicitement promu "
    "en stage Production dans le Registry."
)

h(2, "7.1  Pourquoi un registre de modèles ?")
para(
    "Sans registre, n'importe quel fichier .joblib peut être uploadé et mis en production "
    "par erreur. Avec un registre, le workflow est formalisé :"
)
schema_line("  Entraînement  →  Enregistrement dans le Registry (stage: None)")
schema_line("       ↓         [revue des métriques dans l'UI]")
schema_line("  Promotion en Staging  (tests complémentaires)")
schema_line("       ↓         [validation finale]")
schema_line("  Promotion en Production  (seul ce stage déclenche l'upload)")
schema_line("       ↓         [si remplacé par une version suivante]")
schema_line("  Archived  (conservé mais inactif)")
doc.add_paragraph()

para(
    "Dans ce projet, on passe directement de None à Production "
    "(on saute Staging pour simplifier). Dans une vraie équipe, Staging permet "
    "de faire valider le modèle par un autre data scientist ou un responsable "
    "avant de le mettre face aux utilisateurs."
)

h(2, "7.2  Voir les modèles dans l'UI")
numbered(
    "Dans MLFlow UI (http://127.0.0.1:5000), cliquer sur l'onglet « Models » (en haut)."
)
numbered(
    "La liste affiche : housing_model_A, housing_model_B, housing_latest, housing_v3."
)
numbered(
    "Cliquer sur housing_model_A : on voit la version 1 avec son stage (Production), "
    "le run qui l'a produite, la date de création."
)
numbered(
    "En cliquant sur un numéro de version, on accède aux métriques et paramètres "
    "du run associé — traçabilité complète."
)

h(2, "7.3  Promouvoir les modèles via Python")
para(
    "Dans ce projet, la promotion est faite via un script Python après l'exécution "
    "du notebook. Voici le code exact utilisé :"
)
code(
    "import warnings\n"
    "warnings.filterwarnings('ignore')  # masquer les FutureWarning de déprecation\n"
    "\n"
    "from mlflow.tracking import MlflowClient\n"
    "\n"
    "client = MlflowClient()\n"
    "\n"
    "for model_name in ['housing_model_A', 'housing_model_B']:\n"
    "    # Récupérer toutes les versions et prendre la plus récente\n"
    "    versions = client.search_model_versions(f\"name='{model_name}'\")\n"
    "    latest = max(versions, key=lambda v: int(v.version))\n"
    "\n"
    "    # Promouvoir en Production en archivant les versions précédentes\n"
    "    client.transition_model_version_stage(\n"
    "        name=model_name,\n"
    "        version=latest.version,\n"
    "        stage='Production',\n"
    "        archive_existing_versions=True,\n"
    "    )\n"
    "    print(f'{model_name} v{latest.version} => Production')"
)
tip(
    "ℹ️ FutureWarning sur les stages :",
    "MLFlow 2.9+ affiche un avertissement indiquant que les stages (Staging, Production) "
    "seront dépréciés dans une future version majeure au profit des 'aliases' (@production). "
    "Dans MLFlow 2.22.0, les stages fonctionnent encore parfaitement. Pour un projet "
    "en production longue durée, consulter la documentation de migration MLFlow.",
    fill="D1ECF1",
)

h(2, "7.4  Charger un modèle depuis le Registry")
para(
    "Une fois un modèle en stage Production, on peut le charger par son nom symbolique "
    "sans connaître son run_id ou son chemin physique :"
)
code(
    "import mlflow.sklearn\n"
    "\n"
    "# Charger housing_model_A en stage Production\n"
    "model = mlflow.sklearn.load_model('models:/housing_model_A/Production')\n"
    "\n"
    "# Faire une prédiction\n"
    "import pandas as pd\n"
    "sample = pd.DataFrame([{\n"
    "    'MedInc': 3.5, 'HouseAge': 20, 'AveRooms': 5,\n"
    "    'AveBedrms': 1, 'Population': 1000, 'AveOccup': 3,\n"
    "    'Latitude': 37.0, 'Longitude': -120.0\n"
    "}])\n"
    "prediction = model.predict(sample)\n"
    "print(f'Prix estimé : {prediction[0]:.2f} (x100k$)')"
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# PARTIE 8 — CONNEXION REGISTRY → HF HUB → RENDER
# ═══════════════════════════════════════════════════════════════════════════════
h(1, "Partie 8 — Connexion Registry → HuggingFace Hub → Render")

para(
    "C'est la partie la plus avancée du projet et la plus représentative d'un workflow "
    "MLOps professionnel. Le script scripts/upload_model_to_hf.py est le chaînon "
    "qui relie MLFlow (local) à HuggingFace Hub (stockage cloud) et donc à Render (production)."
)

h(2, "8.1  Pourquoi passer par le Registry pour uploader ?")
para(
    "Dans une version naïve, le script uploadait directement les fichiers .joblib "
    "présents dans artifacts/models/. Cela posait un problème fondamental : "
    "n'importe quel modèle, qu'il soit bon ou mauvais, pouvait partir en production "
    "sans validation."
)
para(
    "Avec la version actuelle, le script refuse d'uploader un modèle qui n'est pas "
    "en stage Production dans le Registry. La promotion devient une étape obligatoire "
    "et tracée dans l'historique MLFlow."
)

h(2, "8.2  Fonctionnement du script upload_model_to_hf.py")
para("Le script réalise trois opérations distinctes :")
numbered(
    "Charger housing_model_A @ Production depuis le Registry "
    "→ le sérialiser en model_v1.joblib dans un dossier temporaire."
)
numbered(
    "Charger housing_model_B @ Production depuis le Registry "
    "→ le sérialiser en model_v2.joblib dans un dossier temporaire."
)
numbered(
    "Chercher le run avec le meilleur R² dans l'expérience housing_price_prediction "
    "→ charger le modèle de ce run → le sérialiser en model_latest.joblib."
)
para("Puis les trois fichiers sont uploadés sur HuggingFace Hub en une seule passe.")

h(2, "8.3  Extraits du script")
code(
    "# Charger un modèle Production depuis le Registry\n"
    "def _load_production_model(client, model_name):\n"
    "    versions = client.get_latest_versions(model_name, stages=['Production'])\n"
    "    if not versions:\n"
    "        raise RuntimeError(\n"
    "            f\"Aucune version en Production pour '{model_name}'.\\n\"\n"
    '            "Promouvoir d\'abord via client.transition_model_version_stage()"\n'
    "        )\n"
    "    return mlflow.sklearn.load_model(f'models:/{model_name}/Production')\n"
    "\n"
    "\n"
    "# Charger le modèle avec le meilleur R²\n"
    "def _load_best_r2_model(client):\n"
    "    experiment = client.get_experiment_by_name('housing_price_prediction')\n"
    "    runs = client.search_runs(\n"
    "        experiment_ids=[experiment.experiment_id],\n"
    "        filter_string='metrics.r2 > 0',\n"
    "        order_by=['metrics.r2 DESC'],\n"
    "        max_results=1,\n"
    "    )\n"
    "    return mlflow.sklearn.load_model(f'runs:/{runs[0].info.run_id}/model')"
)

h(2, "8.4  Workflow complet de mise à jour d'un modèle en production")
para("Voici la procédure complète, de l'entraînement au déploiement :")
numbered(
    "Exécuter le notebook 02 : notebooks/02_train_ab_models.ipynb "
    "(ou python main.py pour le pipeline principal)."
)
numbered(
    "Ouvrir l'UI MLFlow et comparer les métriques des runs : "
    "mlflow ui → http://127.0.0.1:5000"
)
numbered(
    "Décider quel modèle mérite d'aller en production en se basant sur les métriques "
    "(R² le plus élevé, MAE la plus faible)."
)
numbered(
    "Promouvoir le modèle choisi via Python (voir Partie 7.3) ou via l'UI MLFlow "
    "(onglet Models → choisir la version → Stage → Transition to Production)."
)
numbered(
    "Lancer le script d'upload : python scripts/upload_model_to_hf.py\n"
    "Le script charge le modèle depuis le Registry et l'uploade sur HuggingFace Hub."
)
numbered(
    "Le backend Render télécharge automatiquement les nouveaux .joblib au prochain "
    "redémarrage du conteneur (ou déclencher un redeploy sur render.com)."
)
tip(
    "💡 Point clé MLOps :",
    "La promotion en Production n'est pas automatique. C'est un acte humain délibéré "
    "qui constitue la validation du modèle. Cette décision est tracée dans MLFlow "
    "avec une date et une version. C'est ce qu'on appelle le 'human-in-the-loop' "
    "dans un pipeline MLOps.",
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# PARTIE 9 — RÉCAPITULATIF
# ═══════════════════════════════════════════════════════════════════════════════
h(1, "Partie 9 — Récapitulatif des fichiers du projet")

files_summary = [
    (
        "requirements.txt",
        "Ajout de mlflow==2.22.0",
        "Installer MLFlow dans l'environnement de développement local",
    ),
    (
        ".gitignore",
        "Ajout de mlruns/ et mlartifacts/",
        "Ne pas commiter la base de données MLFlow locale (peut dépasser 1 Go)",
    ),
    (
        "src/training/pipeline.py",
        "set_experiment + start_run + log_param/metric/model/artifact",
        "Tracer chaque run du pipeline de production",
    ),
    (
        "notebooks/02_train_ab_models.ipynb",
        "set_tracking_uri + set_experiment + start_run dans cellules 4 et 5",
        "Tracer et enregistrer dans le Registry les modèles A et B",
    ),
    (
        "scripts/upload_model_to_hf.py",
        "Réécriture complète : lit le Registry au lieu de artifacts/",
        "MLFlow devient la porte de production : seul un modèle en Production est uploadé",
    ),
]
table2 = doc.add_table(rows=1, cols=3)
table2.style = "Table Grid"
hdr2 = table2.rows[0].cells
hdr2[0].text = "Fichier"
hdr2[1].text = "Modification"
hdr2[2].text = "Objectif"
for cell in hdr2:
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)
for fname, mod, obj in files_summary:
    rc = table2.add_row().cells
    rc[0].text = fname
    rc[1].text = mod
    rc[2].text = obj
    for c in rc:
        c.paragraphs[0].runs[0].font.size = Pt(9)
doc.add_paragraph()

para("Fichiers NON modifiés (et pourquoi) :")
bullet(
    "backend/requirements.txt — MLFlow est un outil de développement. "
    "Le backend Render ne fait pas de tracking, il charge juste un .joblib."
)
bullet(
    "backend/app.py — Le backend continue de charger les .joblib depuis HF Hub. "
    "Il ne sait pas que MLFlow existe."
)
bullet(
    "src/training/train.py et evaluate.py — La logique métier d'entraînement "
    "reste identique. MLFlow s'intègre autour d'elle sans la modifier."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# ANNEXE
# ═══════════════════════════════════════════════════════════════════════════════
h(1, "Annexe — Aide-mémoire et flux de soutenance")

h(2, "A.1  Fonctions Python essentielles")
api_cmds = [
    (
        "mlflow.set_tracking_uri(uri)",
        "Définir où MLFlow stocke les données (crucial dans les notebooks)",
    ),
    ("mlflow.set_experiment(nom)", "Créer ou sélectionner une expérience par nom"),
    ("mlflow.start_run(run_name=nom)", "Ouvrir un run (utiliser avec le bloc with)"),
    (
        "mlflow.log_param(clé, valeur)",
        "Loguer un paramètre unique (avant entraînement)",
    ),
    ("mlflow.log_params(dict)", "Loguer plusieurs paramètres depuis un dictionnaire"),
    (
        "mlflow.log_metric(clé, valeur)",
        "Loguer une métrique numérique (après entraînement)",
    ),
    ("mlflow.log_metrics(dict)", "Loguer plusieurs métriques depuis un dictionnaire"),
    (
        "mlflow.log_artifact(chemin)",
        "Copier un fichier dans MLFlow (JSON, PNG, CSV...)",
    ),
    (
        "mlflow.sklearn.log_model(model, path, registered_model_name=...)",
        "Sauvegarder un modèle sklearn dans le Registry",
    ),
    (
        "mlflow.sklearn.load_model('models:/nom/Production')",
        "Charger un modèle depuis le Registry",
    ),
    ("mlflow.set_tag(clé, valeur)", "Ajouter une métadonnée libre au run"),
]
for fn, desc in api_cmds:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.3)
    r1 = p.add_run(fn + "\n")
    set_font(r1, bold=True, name="Courier New", size=9, color=(0x1F, 0x49, 0x7D))
    r2 = p.add_run("    → " + desc)
    set_font(r2, size=10)

h(2, "A.2  Méthodes MlflowClient essentielles")
client_cmds = [
    ("client.get_experiment_by_name(nom)", "Récupérer une expérience par son nom"),
    ("client.search_runs([exp_id], order_by=[...])", "Chercher et trier des runs"),
    (
        "client.search_model_versions(f\"name='modele'\")",
        "Lister les versions d'un modèle",
    ),
    (
        "client.get_latest_versions(nom, stages=['Production'])",
        "Récupérer la version en Production",
    ),
    (
        "client.transition_model_version_stage(nom, version, stage)",
        "Promouvoir un modèle",
    ),
]
for fn, desc in client_cmds:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.3)
    r1 = p.add_run(fn + "\n")
    set_font(r1, bold=True, name="Courier New", size=9, color=(0x1F, 0x49, 0x7D))
    r2 = p.add_run("    → " + desc)
    set_font(r2, size=10)

h(2, "A.3  Commandes terminal")
terminal_cmds = [
    ("mlflow ui", "Lancer l'interface sur http://127.0.0.1:5000"),
    ("mlflow ui --port 5001", "Lancer sur un autre port"),
    (
        "python scripts/upload_model_to_hf.py",
        "Uploader les modèles Production vers HF Hub",
    ),
]
for cmd, desc in terminal_cmds:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.3)
    r1 = p.add_run(cmd)
    set_font(r1, bold=True, name="Courier New", size=9, color=(0x1F, 0x49, 0x7D))
    r2 = p.add_run("  →  " + desc)
    set_font(r2, size=10)

h(2, "A.4  Flux recommandé pour la soutenance")
para(
    "Pour démontrer l'intégration MLFlow dans ce projet lors de la soutenance, "
    "suivre ce scénario dans l'ordre :"
)
workflow = [
    "Lancer mlflow ui depuis la racine du projet et ouvrir http://127.0.0.1:5000.",
    "Montrer l'expérience housing_price_prediction avec les runs du pipeline principal "
    "(run_v1, run_v2, run_v3...). Pointer les colonnes Parameters et Metrics.",
    "Passer à l'expérience housing_ab_models. Montrer les deux runs : "
    "model_A_RandomForest et model_B_GradientBoosting.",
    "Sélectionner les deux runs et cliquer sur Compare. Expliquer le tableau comparatif "
    "R²/MAE et justifier pourquoi model_A est le champion.",
    "Cliquer sur l'onglet Models. Montrer housing_model_A et housing_model_B "
    "avec leur stage Production.",
    "Ouvrir housing_model_A → version 1 → montrer le lien vers le run qui l'a produite. "
    "C'est la traçabilité complète : du déploiement à l'expérience.",
    "Expliquer le rôle du script upload_model_to_hf.py : il refuse d'uploader "
    "un modèle qui n'est pas en Production dans le Registry. MLFlow est la porte de production.",
    "Conclure en montrant l'architecture complète : "
    "notebook → MLFlow Registry → upload_model_to_hf.py → HuggingFace Hub → backend Render.",
]
for step in workflow:
    numbered(step)

tip(
    "💡 Message clé à retenir pour la soutenance :",
    "MLFlow dans ce projet ne sert pas qu'à comparer des métriques. "
    "Il formalise la gouvernance : un modèle ne peut pas partir en production "
    "sans une promotion explicite dans le Registry. C'est ce qui distingue "
    "un projet ML amateur d'un projet ML industrialisé.",
)

# Footer
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    "Guide rédigé pour le projet fil rouge — Industrialisation ML  |  MLFlow 2.22.0"
)
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ── Sauvegarde ────────────────────────────────────────────────────────────────
out_path = "d:/Projects/ml-housing-project/docs/Guide_MLFlow.docx"
os.makedirs("d:/Projects/ml-housing-project/docs", exist_ok=True)
doc.save(out_path)
print(f"Sauvegardé : {out_path}")
